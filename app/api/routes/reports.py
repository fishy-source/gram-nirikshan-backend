"""
PDF Report Generator using ReportLab.
Generates department-format inspection reports with photos, GPS data, and approval trail.
"""
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, inch
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    Image as RLImage, HRFlowable, KeepTogether
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from pathlib import Path
from datetime import datetime, timezone
import uuid, os, shutil

try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    docx = None

from app.db.database import get_db
from app.models.models import Inspection, Panchayat, User, Photo, Approval, Report
from app.schemas.schemas import MessageResponse
from app.core.dependencies import get_current_user
from app.core.config import settings

router = APIRouter(prefix="/reports", tags=["Reports"])

REPORTS_DIR = Path(settings.UPLOAD_DIR) / "reports"
REPORTS_DIR.mkdir(parents=True, exist_ok=True)

# Color scheme
PRIMARY = colors.HexColor("#1a5276")
SECONDARY = colors.HexColor("#2e86c1")
ACCENT = colors.HexColor("#f39c12")
LIGHT_BG = colors.HexColor("#eaf2fb")
DARK_BG = colors.HexColor("#154360")


def find_project_root() -> Path:
    current = Path(__file__).resolve().parent
    for _ in range(5):
        if (current / "flutter_app").exists():
            return current
        current = current.parent
    return Path(__file__).parents[3]


def get_absolute_path(rel_path: str) -> Path:
    """Find absolute path of a file, searching workspace directories."""
    path = Path(rel_path)
    if path.exists():
        return path
    root = find_project_root()
    path2 = root / rel_path
    if path2.exists():
        return path2
    path3 = root / "backend" / rel_path
    if path3.exists():
        return path3
    return path


# Register Noto Sans Devanagari Hindi Unicode font
pdf_font_name = 'Helvetica'
pdf_font_bold_name = 'Helvetica-Bold'

try:
    root = find_project_root()
    backend_font_dir = root / "assets" / "fonts"
    backend_font_dir.mkdir(parents=True, exist_ok=True)

    reg_path = backend_font_dir / "NotoSansDevanagari-Regular.ttf"
    bold_path = backend_font_dir / "NotoSansDevanagari-Bold.ttf"
    kd_path = backend_font_dir / "Kruti_Dev_010.ttf"

    import requests
    if not reg_path.exists():
        try:
            r = requests.get('https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSansDevanagari/NotoSansDevanagari-Regular.ttf', timeout=10)
            with open(reg_path, 'wb') as f:
                f.write(r.content)
        except Exception as e:
            print(f"Failed to download Noto Regular font: {e}")
            
    if not bold_path.exists():
        try:
            r = requests.get('https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSansDevanagari/NotoSansDevanagari-Bold.ttf', timeout=10)
            with open(bold_path, 'wb') as f:
                f.write(r.content)
        except Exception as e:
            print(f"Failed to download Noto Bold font: {e}")

    if not kd_path.exists():
        try:
            r = requests.get('https://github.com/ashoksihmar/iascgl/raw/master/Kruti_Dev_010.ttf', timeout=10)
            with open(kd_path, 'wb') as f:
                f.write(r.content)
        except Exception as e:
            print(f"Failed to download Kruti Dev font: {e}")

    if reg_path.exists() and bold_path.exists():
        pdfmetrics.registerFont(TTFont('NotoDev', str(reg_path)))
        pdfmetrics.registerFont(TTFont('NotoDev-Bold', str(bold_path)))
        pdf_font_name = 'NotoDev'
        pdf_font_bold_name = 'NotoDev-Bold'

    if kd_path.exists():
        pdfmetrics.registerFont(TTFont('KrutiDev', str(kd_path)))
        pdfmetrics.registerFont(TTFont('KrutiDev-Bold', str(kd_path)))
except Exception as e:
    print(f"Error registering Noto Sans Devanagari/Kruti Dev fonts: {e}")
    pdf_font_name = 'Helvetica'
    pdf_font_bold_name = 'Helvetica-Bold'


def Unicode_to_KrutiDev(unicode_substring: str) -> str:
    if not unicode_substring:
        return ""
    modified_substring = unicode_substring
    
    array_one = ["‘",   "’",   "“",   "”",   "(",    ")",   "{",    "}",   "=", "।",  "?",  "-",  "µ", "॰", ",", ".", "् ", 
    "०",  "१",  "२",  "३",     "४",   "५",  "६",   "७",   "८",   "९", "x", 
    
    "फ़्",  "क़",  "ख़",  "ग़", "ज़्", "ज़",  "ड़",  "ढ़",   "फ़",  "य़",  "ऱ",  "ऩ",  
    "त्त्",   "त्त",     "क्त",  "दृ",  "कृ",
    
    "ह्न",  "ह्य",  "हृ",  "ह्म",  "ह्र",  "ह्",   "द्द",  "क्ष्", "क्ष", "त्र्", "त्र","ज्ञ",
    "छ्य",  "ट्य",  "ठ्य",  "ड्य",  "ढ्य", "द्य","द्व",
    "श्र",  "ट्र",    "ड्र",    "ढ्र",    "छ्र",   "क्र",  "फ्र",  "द्र",   "प्र",   "ग्र", "रु",  "रू",
    "्र",
    
    "ओ",  "औ",  "आ",   "अ",   "ई",   "इ",  "उ",   "ऊ",  "ऐ",  "ए", "ऋ",
    
    "क्",  "क",  "क्क",  "ख्",   "ख",    "ग्",   "ग",  "घ्",  "घ",    "ङ",
    "चै",   "च्",   "च",   "छ",  "ज्", "ज",   "झ्",  "झ",   "ञ",
    
    "ट्ट",   "ट्ठ",   "ट",   "ठ",   "ड्ड",   "ड्ढ",  "ड",   "ढ",  "ण्", "ण",  
    "त्",  "त",  "थ्", "थ",  "द्ध",  "द", "ध्", "ध",  "न्",  "न",  
    
    "प्",  "प",  "फ्", "फ",  "ब्",  "ब", "भ्",  "भ",  "म्",  "म",
    "य्",  "य",  "र",  "ल्", "ल",  "ळ",  "व्",  "व", 
    "श्", "श",  "ष्", "ष",  "स्",   "स",   "ह",     
    
    "ऑ",   "ॉ",  "ो",   "ौ",   "ा",   "ी",   "ु",   "ू",   "ृ",   "े",   "ै",
    "ं",   "ँ",   "ः",   "ॅ",    "ऽ",  "् ", "्" ]
    
    array_two = ["^", "*",  "Þ", "ß", "¼", "½", "¿", "À", "¾", "A", "\\", "&", "&", "Œ", "]","-","~ ", 
    "å",  "ƒ",  "„",   "…",   "†",   "‡",   "ˆ",   "‰",   "Š",   "‹","Û",
    
    "¶",   "d",    "[k",  "x",  "T",  "t",   "M+", "<+", "Q",  ";",    "j",   "u",
    "Ù",   "Ùk",   "Dr",    "–",   "—",       
    
    "à",   "á",    "â",   "ã",   "ºz",  "º",   "í", "{", "{k",  "«", "=","K", 
    "Nî",   "Vî",    "Bî",   "Mî",   "<î", "|","}",
    "J",   "Vª",   "Mª",  "<ªª",  "Nª",   "Ø",  "Ý",   "æ", "ç", "xz", "#", ":",
    "z",
    
    "vks",  "vkS",  "vk",    "v",   "bZ",  "b",  "m",  "Å",  ",s",  ",",   "_",
    
    "D",  "d",    "ô",     "[",     "[k",    "X",   "x",  "?",    "?k",   "³", 
    "pkS",  "P",    "p",  "N",   "T",    "t",   "÷",  ">",   "¥",
    
    "ê",      "ë",      "V",  "B",   "ì",       "ï",     "M",  "<",  ".", ".k",   
    "R",  "r",   "F", "Fk",  ")",    "n", "/",  "/k",  "U", "u",   
    
    "I",  "i",   "¶", "Q",   "C",  "c",  "H",  "Hk", "E",   "e",
    "¸",   ";",    "j",  "Y",   "y",  "G",  "O",  "o",
    "'", "'k",  "\"", "\"k", "L",   "l",   "g",      
    
    "v‚",    "‚",    "ks",   "kS",   "k",     "h",    "q",   "w",   "`",    "s",    "S",
    "a",    "¡",    "%",     "W",   "·",   "~ ", "~"]
    
    array_one_length = len(array_one)
    
    modified_substring = modified_substring.replace("ि", "ि")
    modified_substring = modified_substring.replace("ि", "f")
    
    for input_symbol_idx in range(0, array_one_length):
        modified_substring = modified_substring.replace(array_one[input_symbol_idx], array_two[input_symbol_idx])
    
    # Move "f" to correct position (before consonant)
    modified_substring = "  " + modified_substring + "  "
    position_of_f = modified_substring.find("f")
    while position_of_f != -1:
        modified_substring = modified_substring[:position_of_f-1] + "f" + modified_substring[position_of_f-1] + modified_substring[position_of_f+1:]
        position_of_f = modified_substring.find("f", position_of_f + 1)
    modified_substring = modified_substring.strip()
    
    # Move "half R" to correct position
    modified_substring = "  " + modified_substring + "  "
    position_of_r = modified_substring.find("j~")
    set_of_matras = ["‚", "ks", "kS", "k", "h", "q", "w", "`", "s", "S", "a", "¡", "%", "W", "·", "~ ", "~"]
    while position_of_r != -1:
        modified_substring = modified_substring.replace("j~", "", 1)
        if modified_substring[position_of_r + 1] in set_of_matras:
            modified_substring = modified_substring[:position_of_r + 2] + "Z" + modified_substring[position_of_r + 2:]
        else:
            modified_substring = modified_substring[:position_of_r + 1] + "Z" + modified_substring[position_of_r + 1:]
        position_of_r = modified_substring.find("j~")
    modified_substring = modified_substring.strip()
    
    return modified_substring


def to_pdf_html(text: str, bold: bool = False, size: float = None) -> str:
    if not text:
        return ""
    spans = []
    current_type = None
    current_str = []
    for c in text:
        # Check if character is Hindi or ZWJ/ZWNJ
        is_hindi = (0x0900 <= ord(c) <= 0x097F) or (ord(c) in (0x200c, 0x200d))
        if is_hindi:
            this_type = 'hindi'
        elif c in (' ', ',', '.', ':', '(', ')', '-', '/', '_', '[', ']', '\n', '\r'):
            this_type = current_type if current_type else 'english'
        else:
            this_type = 'english'
            
        if this_type != current_type:
            if current_str:
                spans.append((current_type, ''.join(current_str)))
            current_type = this_type
            current_str = [c]
        else:
            current_str.append(c)
    if current_str:
        spans.append((current_type, ''.join(current_str)))

    html_parts = []
    for t, s in spans:
        if t == 'hindi':
            # Convert Unicode to KrutiDev for correct rendering in ReportLab
            kruti_s = Unicode_to_KrutiDev(s)
            escaped_s = kruti_s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('\n', '<br/>').replace('\r', '')
            font = 'KrutiDev-Bold' if bold else 'KrutiDev'
            # Kruti Dev typically needs to be slightly larger than standard fonts to be readable
            fs = f' size="{size + 2.5}"' if size else ' size="12"'
            html_parts.append(f'<font name="{font}"{fs}>{escaped_s}</font>')
        else:
            escaped_s = s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('\n', '<br/>').replace('\r', '')
            font = 'Helvetica-Bold' if bold else 'Helvetica'
            fs = f' size="{size}"' if size else ''
            html_parts.append(f'<font name="{font}"{fs}>{escaped_s}</font>')
            
    return "".join(html_parts)


def extract_witness_name(description: str) -> str:
    if not description:
        return "___________________"
    import re
    match = re.search(r"(?:गवाह का नाम|गवाह\s*नाम|गवाह|witness name|witness)\s*[:：\-‐—\s]\s*([^\n\r।\.]+)", description, re.IGNORECASE)
    if match:
        name = match.group(1).strip()
        name = name.split(',')[0].split(';')[0].strip()
        name = re.sub(r"^(?:श्री|shri|mr\.|mrs\.)\s+", "", name, flags=re.IGNORECASE)
        stop_words = ["उपस्थित", "मौजूद", "थे", "था", "रहा", "रही", "गया", "गई", "है", "हैं", "present", "here", "was", "is", "were"]
        for stop_word in stop_words:
            name = re.sub(rf"(?:\s+|^){stop_word}(?:\s+|$)", " ", name, flags=re.IGNORECASE).strip()
        if name:
            return name
    return "___________________"


def get_val_str(val) -> str:
    if val is None:
        return ""
    if hasattr(val, "value"):
        return str(val.value)
    return str(val)


def get_formatted_district_block(inspection, panchayat):
    dist = inspection.district or (panchayat.district if panchayat else "N/A")
    dist_clean = "".join(dist.lower().split()) if dist else ""
    if "hathras" in dist_clean:
        dist = "Hathras"
    elif "sikandra" in dist_clean:
        dist = "Sikandra Rao"
    
    blk = inspection.block or (panchayat.block if panchayat else "N/A")
    blk_clean = "".join(blk.lower().split()) if blk else ""
    if "hathras" in blk_clean:
        blk = "Hathras"
    elif "sikandra" in blk_clean:
        blk = "Sikandra Rao"
        
    return dist, blk


def build_docx_report(inspection, panchayat, engineer, photos, approvals, output_path: str):
    """Build a formal Microsoft Word report in Kruti Dev 010."""
    if not docx:
        raise RuntimeError("python-docx library is not installed on this server.")
        
    has_hindi = any(v and isinstance(v, str) and any(0x0900 <= ord(c) <= 0x097F for c in v) 
                    for v in [inspection.title, inspection.project_name, inspection.observations, inspection.recommendations, inspection.ai_report_draft])

    def L(en_label: str, hi_label: str) -> str:
        return f"{hi_label} / {en_label}" if has_hindi else en_label

    doc = docx.Document()
    
    # Set margins to 0.75 inches for page-like structure
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Add a professional page border using XML manipulation
    try:
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        for section in doc.sections:
            sectPr = section._sectPr
            pgBorders_xml = (
                f'<w:pgBorders {nsdecls("w")}>\n'
                f'  <w:top w:val="single" w:sz="6" w:space="24" w:color="1A5276"/>\n'
                f'  <w:left w:val="single" w:sz="6" w:space="24" w:color="1A5276"/>\n'
                f'  <w:bottom w:val="single" w:sz="6" w:space="24" w:color="1A5276"/>\n'
                f'  <w:right w:val="single" w:sz="6" w:space="24" w:color="1A5276"/>\n'
                f'</w:pgBorders>'
            )
            sectPr.append(parse_xml(pgBorders_xml))
    except Exception as border_err:
        print(f"Failed to add page borders: {border_err}")

    # Set document header/footer
    try:
        header = doc.sections[0].header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_title = L("Gram Panchayat Inspection Report", "ग्राम पंचायत निरीक्षण रिपोर्ट")
        hrun = hp.add_run(Unicode_to_KrutiDev(header_title) if has_hindi else header_title)
        if has_hindi:
            hrun.font.name = 'Kruti Dev 010'
            hrun.font.size = Pt(14)
        else:
            hrun.font.name = 'Calibri'
            hrun.font.size = Pt(8.5)
        hrun.font.italic = not has_hindi
        hrun.font.color.rgb = RGBColor(128, 128, 128)
    except Exception as header_err:
        print(f"Failed to add header: {header_err}")

    # Helper function to classify and write mixed text runs in Kruti Dev or Calibri
    def add_runs_to_p(p, text, bold=False, font_size_pt=11, color_rgb=None):
        if not text:
            return
        spans = []
        current_type = None
        current_str = []
        for c in text:
            is_hindi = (0x0900 <= ord(c) <= 0x097F)
            if is_hindi:
                this_type = 'hindi'
            elif c in (' ', ',', '.', ':', '(', ')'):
                this_type = current_type if current_type else 'english'
            else:
                this_type = 'english'
            if this_type != current_type:
                if current_str:
                    spans.append((current_type, ''.join(current_str)))
                current_type = this_type
                current_str = [c]
            else:
                current_str.append(c)
        if current_str:
            spans.append((current_type, ''.join(current_str)))

        for t, s in spans:
            run = p.add_run()
            if t == 'hindi':
                run.text = Unicode_to_KrutiDev(s)
                run.font.name = 'Kruti Dev 010'
                run.font.size = Pt(font_size_pt + 6)
            else:
                run.text = s
                run.font.name = 'Calibri'
                run.font.size = Pt(font_size_pt)
            if bold:
                run.bold = True
            if color_rgb:
                run.font.color.rgb = RGBColor(*color_rgb)

    def set_cell_background(cell, color_hex):
        try:
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
            cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))
        except Exception as e:
            print(f"Failed to set background color: {e}")

    def set_cell(cell, text, bold=False, font_size_pt=10.5, color_rgb=None, is_label=False):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        add_runs_to_p(p, text, bold=bold, font_size_pt=font_size_pt, color_rgb=color_rgb)
        if is_label:
            set_cell_background(cell, "EAF2FB")

    def set_table_widths(table, widths):
        for row in table.rows:
            for i, w in enumerate(widths):
                if i < len(row.cells):
                    row.cells[i].width = w

    # Document Header Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(18)
    add_runs_to_p(p_title, L("INSPECTION REPORT", "निरीक्षण रिपोर्ट"), bold=True, font_size_pt=24, color_rgb=(26, 82, 118))

    # Basic Information
    engineer_name = inspection.investigator_name or (engineer.name_hindi or engineer.name if engineer else "N/A")
    dist, blk = get_formatted_district_block(inspection, panchayat)
    
    # Subheading/Section: Basic Info
    p_info_lbl = doc.add_paragraph()
    p_info_lbl.paragraph_format.space_before = Pt(12)
    p_info_lbl.paragraph_format.space_after = Pt(6)
    p_info_lbl.paragraph_format.keep_with_next = True
    add_runs_to_p(p_info_lbl, L("Inspection Basic Details", "निरीक्षण के बुनियादी विवरण"), bold=True, font_size_pt=14, color_rgb=(26, 82, 118))
    
    basic_table = doc.add_table(rows=5, cols=4)
    basic_table.style = 'Table Grid'
    set_table_widths(basic_table, [Inches(1.5), Inches(2.0), Inches(1.5), Inches(2.0)])
    
    status_val = get_val_str(inspection.status).lower()
    status_labels = {
        "draft": L("Draft", "प्रारूप (Draft)"),
        "submitted": L("Submitted", "प्रस्तुत (Submitted)"),
        "verified": L("Verified", "सत्यापित (Verified)"),
        "approved": L("Approved", "स्वीकृत (Approved)"),
        "rejected": L("Rejected", "अस्वीकृत (Rejected)")
    }
    status_eng = status_labels.get(status_val, status_val.upper())

    # Row 0
    set_cell(basic_table.rows[0].cells[0], L("Inspection ID", "निरीक्षण आईडी"), bold=True, is_label=True, color_rgb=(26, 82, 118))
    set_cell(basic_table.rows[0].cells[1], inspection.inspection_id)
    set_cell(basic_table.rows[0].cells[2], L("Status", "स्थिति"), bold=True, is_label=True, color_rgb=(26, 82, 118))
    set_cell(basic_table.rows[0].cells[3], status_eng, bold=True)
    
    # Row 1
    set_cell(basic_table.rows[1].cells[0], L("Inspector", "निरीक्षक"), bold=True, is_label=True, color_rgb=(26, 82, 118))
    set_cell(basic_table.rows[1].cells[1], engineer_name)
    set_cell(basic_table.rows[1].cells[2], L("District", "जनपद"), bold=True, is_label=True, color_rgb=(26, 82, 118))
    set_cell(basic_table.rows[1].cells[3], dist)
    
    # Row 2
    set_cell(basic_table.rows[2].cells[0], L("Gram Panchayat", "ग्राम पंचायत"), bold=True, is_label=True, color_rgb=(26, 82, 118))
    set_cell(basic_table.rows[2].cells[1], panchayat.name_hindi or panchayat.name if panchayat else "N/A")
    set_cell(basic_table.rows[2].cells[2], L("Village", "ग्राम"), bold=True, is_label=True, color_rgb=(26, 82, 118))
    set_cell(basic_table.rows[2].cells[3], panchayat.village or "N/A" if panchayat else "N/A")
    
    # Row 3
    set_cell(basic_table.rows[3].cells[0], L("Development Block", "विकास खंड"), bold=True, is_label=True, color_rgb=(26, 82, 118))
    set_cell(basic_table.rows[3].cells[1], blk)
    set_cell(basic_table.rows[3].cells[2], L("Inspection Date", "निरीक्षण तिथि"), bold=True, is_label=True, color_rgb=(26, 82, 118))
    set_cell(basic_table.rows[3].cells[3], str(inspection.inspection_date)[:10] if inspection.inspection_date else "N/A")
    
    # Row 4
    set_cell(basic_table.rows[4].cells[0], L("Project Name", "परियोजना का नाम"), bold=True, is_label=True, color_rgb=(26, 82, 118))
    set_cell(basic_table.rows[4].cells[1], inspection.project_name or "N/A")
    set_cell(basic_table.rows[4].cells[2], L("Work Code", "कार्य कोड"), bold=True, is_label=True, color_rgb=(26, 82, 118))
    set_cell(basic_table.rows[4].cells[3], inspection.project_code or "N/A")
    
    # GPS Details
    if inspection.checkin_latitude:
        p_gps_lbl = doc.add_paragraph()
        p_gps_lbl.paragraph_format.space_before = Pt(16)
        p_gps_lbl.paragraph_format.space_after = Pt(6)
        p_gps_lbl.paragraph_format.keep_with_next = True
        add_runs_to_p(p_gps_lbl, L("GPS Check-in / Check-out Details", "जीपीएस चेक-इन / चेक-आउट विवरण"), bold=True, font_size_pt=14, color_rgb=(26, 82, 118))
        
        gps_table = doc.add_table(rows=3, cols=4)
        gps_table.style = 'Table Grid'
        set_table_widths(gps_table, [Inches(1.5), Inches(2.0), Inches(1.5), Inches(2.0)])
        
        # Row 0
        set_cell(gps_table.rows[0].cells[0], L("Check-in Time", "चेक-इन समय"), bold=True, is_label=True, color_rgb=(26, 82, 118))
        set_cell(gps_table.rows[0].cells[1], str(inspection.checkin_time)[:16] if inspection.checkin_time else "N/A")
        set_cell(gps_table.rows[0].cells[2], L("Check-in GPS", "चेक-इन जीपीएस"), bold=True, is_label=True, color_rgb=(26, 82, 118))
        set_cell(gps_table.rows[0].cells[3], f"{inspection.checkin_latitude:.6f}, {inspection.checkin_longitude:.6f}")
        
        # Row 1
        set_cell(gps_table.rows[1].cells[0], L("Check-out Time", "चेक-आउट समय"), bold=True, is_label=True, color_rgb=(26, 82, 118))
        set_cell(gps_table.rows[1].cells[1], str(inspection.checkout_time)[:16] if inspection.checkout_time else "N/A")
        set_cell(gps_table.rows[1].cells[2], L("Check-out GPS", "चेक-आउट जीपीएस"), bold=True, is_label=True, color_rgb=(26, 82, 118))
        set_cell(gps_table.rows[1].cells[3], f"{inspection.checkout_latitude:.6f}, {inspection.checkout_longitude:.6f}" if inspection.checkout_latitude else "N/A")
        
        # Row 2 (merged cell for address)
        set_cell(gps_table.rows[2].cells[0], L("Check-in Location", "चेक-इन स्थान"), bold=True, is_label=True, color_rgb=(26, 82, 118))
        
        cell_merged = gps_table.rows[2].cells[1].merge(gps_table.rows[2].cells[2]).merge(gps_table.rows[2].cells[3])
        set_cell(cell_merged, inspection.checkin_address or "N/A")
        
    # Map Attachment
    if inspection.map_image_path:
        map_file_path = get_absolute_path(inspection.map_image_path)
        if map_file_path.exists():
            p_map = doc.add_paragraph()
            p_map.paragraph_format.space_before = Pt(16)
            p_map.paragraph_format.space_after = Pt(6)
            p_map.paragraph_format.keep_with_next = True
            add_runs_to_p(p_map, L("Inspection Location Map", "निरीक्षण स्थान का नक्शा"), bold=True, font_size_pt=14, color_rgb=(26, 82, 118))
            
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                p_img.add_run().add_picture(str(map_file_path), width=Inches(5.5))
            except Exception as e:
                p_img.add_run(f"Failed to load map image: {e}")
            
    # Observations & Recommendations
    for section_title, hi_section_title, content in [
        ("Key Observations / Deficiencies", "मुख्य अवलोकन / कमियां", inspection.observations),
        ("Corrective Recommendations / Measures", "सुधारात्मक सिफारिशें / उपाय", inspection.recommendations),
        ("Action Taken / Remarks", "की गई कार्रवाई / टिप्पणी", inspection.action_taken),
    ]:
        if content:
            p_lbl = doc.add_paragraph()
            p_lbl.paragraph_format.space_before = Pt(16)
            p_lbl.paragraph_format.space_after = Pt(6)
            p_lbl.paragraph_format.keep_with_next = True
            add_runs_to_p(p_lbl, L(section_title, hi_section_title), bold=True, font_size_pt=14, color_rgb=(26, 82, 118))
            
            p_val = doc.add_paragraph()
            p_val.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_val.paragraph_format.space_before = Pt(4)
            p_val.paragraph_format.space_after = Pt(12)
            p_val.paragraph_format.line_spacing = 1.15
            add_runs_to_p(p_val, content, font_size_pt=11)
            
    # AI Report Draft
    if inspection.ai_report_draft:
        p_ai_lbl = doc.add_paragraph()
        p_ai_lbl.paragraph_format.space_before = Pt(16)
        p_ai_lbl.paragraph_format.space_after = Pt(6)
        p_ai_lbl.paragraph_format.keep_with_next = True
        add_runs_to_p(p_ai_lbl, L("Detailed Inspection Report", "विस्तृत निरीक्षण रिपोर्ट"), bold=True, font_size_pt=14, color_rgb=(26, 82, 118))
        
        ai_lines = inspection.ai_report_draft.split('\n')
        for line in ai_lines:
            if line.strip():
                is_bold = line.strip().startswith('**') or line.strip().startswith('###') or line.strip().startswith('##') or line.strip().startswith('*')
                clean_line = line.strip().replace('**', '').replace('###', '').replace('##', '').replace('*', '').strip()
                p_line = doc.add_paragraph()
                p_line.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_line.paragraph_format.space_before = Pt(2)
                p_line.paragraph_format.space_after = Pt(4)
                p_line.paragraph_format.line_spacing = 1.15
                add_runs_to_p(p_line, clean_line, bold=is_bold, font_size_pt=11)
            else:
                p_space = doc.add_paragraph()
                p_space.paragraph_format.space_after = Pt(6)
            
    # Photos
    valid_photos = []
    for p in photos:
        if p.file_path:
            abs_p = get_absolute_path(p.file_path)
            if abs_p.exists():
                valid_photos.append((p, abs_p))
                
    if valid_photos:
        p_photo_lbl = doc.add_paragraph()
        p_photo_lbl.paragraph_format.space_before = Pt(16)
        p_photo_lbl.paragraph_format.space_after = Pt(6)
        p_photo_lbl.paragraph_format.keep_with_next = True
        add_runs_to_p(p_photo_lbl, L("Inspection Site Photographs", "निरीक्षण स्थल के चित्र"), bold=True, font_size_pt=14, color_rgb=(26, 82, 118))
        
        photo_table = doc.add_table(rows=(len(valid_photos) + 1) // 2, cols=2)
        photo_table.style = 'Table Grid'
        set_table_widths(photo_table, [Inches(3.5), Inches(3.5)])
        
        for idx, (photo, abs_p) in enumerate(valid_photos):
            row_idx = idx // 2
            col_idx = idx % 2
            cell = photo_table.rows[row_idx].cells[col_idx]
            cell.text = ""
            p_cell = cell.paragraphs[0]
            p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cell.paragraph_format.space_before = Pt(6)
            p_cell.paragraph_format.space_after = Pt(6)
            try:
                p_cell.add_run().add_picture(str(abs_p), width=Inches(3.0))
                caption_text = f"{photo.caption or 'Site Photograph'} ({str(photo.captured_at)[:16] if photo.captured_at else ''})"
                p_cap = cell.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_before = Pt(2)
                p_cap.paragraph_format.space_after = Pt(4)
                add_runs_to_p(p_cap, caption_text, font_size_pt=9.5, color_rgb=(85, 85, 85))
            except Exception as e:
                p_cell.add_run(f"Failed to load photograph: {e}")
                
    # Approval Details
    if approvals:
        p_app_lbl = doc.add_paragraph()
        p_app_lbl.paragraph_format.space_before = Pt(16)
        p_app_lbl.paragraph_format.space_after = Pt(6)
        p_app_lbl.paragraph_format.keep_with_next = True
        add_runs_to_p(p_app_lbl, L("Approval Details", "स्वीकृति विवरण"), bold=True, font_size_pt=14, color_rgb=(26, 82, 118))
        
        app_table = doc.add_table(rows=1 + len(approvals), cols=5)
        app_table.style = 'Table Grid'
        set_table_widths(app_table, [Inches(1.0), Inches(2.0), Inches(1.2), Inches(1.5), Inches(1.3)])
        
        # Header Row
        headers = [L("Level", "स्तर"), L("Officer", "अधिकारी"), L("Status", "स्थिति"), L("Remarks", "टिप्पणी"), L("Date", "दिनांक")]
        for idx, h in enumerate(headers):
            set_cell(app_table.rows[0].cells[idx], h, bold=True, is_label=True, color_rgb=(26, 82, 118))
            
        for row_idx, a in enumerate(approvals, start=1):
            app_name = a.approver.name_hindi or a.approver.name if a.approver else "N/A"
            desig = a.approver.designation or "Officer" if a.approver else ""
            act_labels = {"pending": L("Pending", "लंबित (Pending)"), "approved": L("Approved", "स्वीकृत (Approved)"), "rejected": L("Rejected", "अस्वीकृत (Rejected)"), "forwarded": L("Forwarded", "प्रेषित (Forwarded)")}
            action_val = get_val_str(a.action)
            action_eng = act_labels.get(action_val.lower(), action_val.upper())
            
            set_cell(app_table.rows[row_idx].cells[0], a.level)
            set_cell(app_table.rows[row_idx].cells[1], f"{app_name} ({desig})")
            set_cell(app_table.rows[row_idx].cells[2], action_eng, bold=True)
            set_cell(app_table.rows[row_idx].cells[3], a.remarks or "-")
            set_cell(app_table.rows[row_idx].cells[4], str(a.created_at)[:16])

    # Signatures
    witness_name = extract_witness_name(inspection.action_taken)
    if witness_name == "___________________":
        witness_name = extract_witness_name(inspection.description)
    
    p_sig_space = doc.add_paragraph()
    p_sig_space.paragraph_format.space_before = Pt(24)
    
    sig_table = doc.add_table(rows=4, cols=3)
    sig_table.style = 'Normal Table'
    set_table_widths(sig_table, [Inches(3.2), Inches(0.6), Inches(3.2)])
    
    set_cell(sig_table.rows[0].cells[0], L("Signature of Inspecting Officer", "निरीक्षण अधिकारी के हस्ताक्षर"), bold=True, font_size_pt=11, color_rgb=(26, 82, 118))
    set_cell(sig_table.rows[0].cells[2], L("Signature of Witness / Representative", "गवाह / प्रतिनिधि के हस्ताक्षर"), bold=True, font_size_pt=11, color_rgb=(26, 82, 118))
    
    set_cell(sig_table.rows[1].cells[0], L(f"Name: {engineer_name}", f"नाम: {engineer_name}"), font_size_pt=10.5)
    set_cell(sig_table.rows[1].cells[2], L(f"Name: {witness_name}", f"नाम: {witness_name}"), font_size_pt=10.5)
    
    set_cell(sig_table.rows[2].cells[0], L(f"Designation: {engineer.designation or 'Junior Engineer' if engineer else 'N/A'}", f"पदनाम: {engineer.designation or 'Junior Engineer' if engineer else 'N/A'}"), font_size_pt=10.5)
    set_cell(sig_table.rows[2].cells[2], "", font_size_pt=10.5)
    
    set_cell(sig_table.rows[3].cells[0], L(f"Date: {datetime.now().strftime('%d/%m/%Y')}", f"दिनांक: {datetime.now().strftime('%d/%m/%Y')}"), font_size_pt=10.5)
    set_cell(sig_table.rows[3].cells[2], L("Date: ___________________", "दिनांक: ___________________"), font_size_pt=10.5)
    
    # Footer
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.space_before = Pt(24)
    p_foot.paragraph_format.space_after = Pt(6)
    footer_text = f"Auto-generated by Gram Nirikshan Mobile App | {datetime.now().strftime('%d/%m/%Y %H:%M')} | Inspection ID: {inspection.inspection_id}"
    if has_hindi:
        footer_text = f"ग्राम निरीक्षण मोबाइल ऐप द्वारा स्वचालित रूप से उत्पन्न | {datetime.now().strftime('%d/%m/%Y %H:%M')} | निरीक्षण आईडी: {inspection.inspection_id}"
    add_runs_to_p(p_foot, f"──────────────────────────────────────────────────\n{footer_text}", font_size_pt=9.5, color_rgb=(128, 128, 128))
    
    doc.save(output_path)


def build_pdf_report(inspection, panchayat, engineer, photos, approvals, output_path: str):
    """Build a simplified PDF inspection report without squares."""
    # Determine language dynamically based on input fields containing Hindi characters
    has_hindi = any(v and isinstance(v, str) and any(0x0900 <= ord(c) <= 0x097F for c in v) 
                    for v in [inspection.title, inspection.project_name, inspection.observations, inspection.recommendations, inspection.ai_report_draft])

    def L(en_label: str, hi_label: str) -> str:
        return f"{hi_label} / {en_label}" if has_hindi else en_label

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=1.5*cm, bottomMargin=2*cm,
    )

    styles = getSampleStyleSheet()
    story = []

    # ── Header ────────────────────────────────────────────────
    title_26 = ParagraphStyle(
        "Title26", fontSize=26, textColor=PRIMARY,
        alignment=TA_CENTER, fontName='Helvetica-Bold', spaceAfter=12
    )
    title_22 = ParagraphStyle(
        "Title22", fontSize=22, textColor=SECONDARY,
        alignment=TA_CENTER, fontName='Helvetica-Bold', spaceAfter=2
    )
    title_18 = ParagraphStyle(
        "Title18", fontSize=18, textColor=colors.black,
        alignment=TA_CENTER, fontName='Helvetica-Bold', spaceAfter=8
    )
    normal = ParagraphStyle("Normal2", fontSize=9, fontName='Helvetica', leading=12, alignment=TA_JUSTIFY, spaceAfter=4)
    label = ParagraphStyle("Label", fontSize=9, fontName='Helvetica-Bold', textColor=PRIMARY, spaceBefore=4, spaceAfter=2)
    label_large = ParagraphStyle("LabelLarge", fontSize=20, fontName='Helvetica-Bold', textColor=PRIMARY, spaceBefore=12, spaceAfter=6, leading=24)

    story.append(Paragraph(to_pdf_html(L("INSPECTION REPORT", "निरीक्षण रिपोर्ट"), bold=True, size=26), title_26))
    story.append(HRFlowable(width="100%", thickness=1.5, color=PRIMARY, spaceBefore=2, spaceAfter=6))

    # ── Basic Information ──────────────────────────────────────
    engineer_name = inspection.investigator_name or (engineer.name_hindi or engineer.name if engineer else "N/A")
    dist, blk = get_formatted_district_block(inspection, panchayat)

    status_val = get_val_str(inspection.status).lower()
    status_labels = {
        "draft": L("Draft", "प्रारूप (Draft)"),
        "submitted": L("Submitted", "प्रस्तुत (Submitted)"),
        "verified": L("Verified", "सत्यापित (Verified)"),
        "approved": L("Approved", "स्वीकृत (Approved)"),
        "rejected": L("Rejected", "अस्वीकृत (Rejected)")
    }
    status_eng = status_labels.get(status_val, status_val.upper())

    info_data = [
        [Paragraph(to_pdf_html(L("Inspection ID", "निरीक्षण आईडी")), normal), Paragraph(to_pdf_html(inspection.inspection_id), normal), 
         Paragraph(to_pdf_html(L("Status", "स्थिति")), normal), Paragraph(to_pdf_html(status_eng), normal)],
        [Paragraph(to_pdf_html(L("Inspector", "निरीक्षक")), normal), Paragraph(to_pdf_html(engineer_name), normal), 
         Paragraph(to_pdf_html(L("District", "जनपद")), normal), Paragraph(to_pdf_html(dist), normal)],
        [Paragraph(to_pdf_html(L("Gram Panchayat", "ग्राम पंचायत")), normal), Paragraph(to_pdf_html(panchayat.name_hindi or panchayat.name if panchayat else "N/A"), normal), 
         Paragraph(to_pdf_html(L("Village", "ग्राम")), normal), Paragraph(to_pdf_html(panchayat.village or "N/A" if panchayat else "N/A"), normal)],
        [Paragraph(to_pdf_html(L("Block", "विकास खंड")), normal), Paragraph(to_pdf_html(blk), normal), 
         Paragraph(to_pdf_html(L("Inspection Date", "निरीक्षण तिथि")), normal), Paragraph(to_pdf_html(str(inspection.inspection_date)[:10] if inspection.inspection_date else "N/A"), normal)],
        [Paragraph(to_pdf_html(L("Project Name", "परियोजना का नाम")), normal), Paragraph(to_pdf_html(inspection.project_name or "N/A"), normal), 
         Paragraph(to_pdf_html(L("Work Code", "कार्य कोड")), normal), Paragraph(to_pdf_html(inspection.project_code or "N/A"), normal)],
    ]

    info_table = Table(info_data, colWidths=[4*cm, 5.5*cm, 4*cm, 5.5*cm])
    info_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), LIGHT_BG),
        ("BACKGROUND", (2, 0), (2, -1), LIGHT_BG),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("PADDING", (0, 0), (-1, -1), 4),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 0.15*cm))

    # ── GPS Information ────────────────────────────────────────
    if inspection.checkin_latitude:
        story.append(Paragraph(to_pdf_html(L("GPS Check-in / Check-out Details", "जीपीएस चेक-इन / चेक-आउट विवरण"), bold=True), label))
        gps_data = [
            [Paragraph(to_pdf_html(L("Check-in Time", "चेक-इन समय")), normal), Paragraph(to_pdf_html(str(inspection.checkin_time)[:16] if inspection.checkin_time else "N/A"), normal),
             Paragraph(to_pdf_html(L("Check-in GPS", "चेक-इन जीपीएस")), normal), Paragraph(to_pdf_html(f"{inspection.checkin_latitude:.6f}, {inspection.checkin_longitude:.6f}"), normal)],
            [Paragraph(to_pdf_html(L("Check-out Time", "चेक-आउट समय")), normal), Paragraph(to_pdf_html(str(inspection.checkout_time)[:16] if inspection.checkout_time else "N/A"), normal),
             Paragraph(to_pdf_html(L("Check-out GPS", "चेक-आउट जीपीएस")), normal), Paragraph(to_pdf_html(f"{inspection.checkout_latitude:.6f}, {inspection.checkout_longitude:.6f}" if inspection.checkout_latitude else "N/A"), normal)],
            [Paragraph(to_pdf_html(L("Check-in Location", "चेक-इन स्थान")), normal), Paragraph(to_pdf_html(inspection.checkin_address or "N/A"), normal), "", ""],
        ]
        gps_table = Table(gps_data, colWidths=[4*cm, 5.5*cm, 4*cm, 5.5*cm])
        gps_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), LIGHT_BG),
            ("BACKGROUND", (2, 0), (2, -1), LIGHT_BG),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("PADDING", (0, 0), (-1, -1), 4),
            ("SPAN", (1, 2), (3, 2)),
        ]))
        story.append(gps_table)
        story.append(Spacer(1, 0.15*cm))

    # ── Map Attachment ─────────────────────────────────────────
    if inspection.map_image_path:
        map_file_path = get_absolute_path(inspection.map_image_path)
        if map_file_path.exists():
            story.append(Paragraph(to_pdf_html(L("Inspection Location Map", "निरीक्षण स्थान का नक्शा"), bold=True), label))
            story.append(Spacer(1, 0.1*cm))
            try:
                img = RLImage(str(map_file_path), width=15*cm, height=6.5*cm)
                story.append(KeepTogether([img, Spacer(1, 0.15*cm)]))
            except Exception as e:
                story.append(Paragraph(to_pdf_html(f"Failed to load map image: {str(e)}"), normal))

    # ── Observations & Recommendations ─────────────────────────
    for section_title, hi_section_title, content in [
        ("Key Observations / Deficiencies", "मुख्य अवलोकन / कमियां", inspection.observations),
        ("Corrective Recommendations / Measures", "सुधारात्मक सिफारिशें / उपाय", inspection.recommendations),
        ("Action Taken / Remarks", "की गई कार्रवाई / टिप्पणी", inspection.action_taken),
    ]:
        if content:
            display_title = L(section_title, hi_section_title)
            story.append(Paragraph(to_pdf_html(display_title + " :-", bold=True, size=20), label_large))
            story.append(Paragraph(to_pdf_html(content), normal))
            story.append(Spacer(1, 0.15*cm))

    # ── AI Report Draft ────────────────────────────────────────
    if inspection.ai_report_draft:
        story.append(Paragraph(to_pdf_html(L("Detailed Inspection Report :-", "विस्तृत निरीक्षण रिपोर्ट :-"), bold=True, size=20), label_large))
        story.append(Spacer(1, 0.1*cm))
        ai_lines = inspection.ai_report_draft.split('\n')
        for line in ai_lines:
            if line.strip():
                # Check if it looks like a markdown heading
                is_bold = line.strip().startswith('**') or line.strip().startswith('###') or line.strip().startswith('##') or line.strip().startswith('*')
                clean_line = line.strip().replace('**', '').replace('###', '').replace('##', '').replace('*', '').strip()
                story.append(Paragraph(to_pdf_html(clean_line, bold=is_bold), normal))
            else:
                story.append(Spacer(1, 0.05*cm))
        story.append(Spacer(1, 0.15*cm))

    # ── Photos ────────────────────────────────────────────────
    valid_photos = []
    for p in photos:
        if p.file_path:
            abs_p = get_absolute_path(p.file_path)
            if abs_p.exists():
                valid_photos.append((p, abs_p))

    if valid_photos:
        story.append(HRFlowable(width="100%", thickness=1, color=SECONDARY, spaceBefore=8, spaceAfter=8))
        story.append(Paragraph(to_pdf_html(L("Inspection Site Photographs :-", "निरीक्षण स्थल के चित्र :-"), bold=True, size=20), label_large))
        story.append(Spacer(1, 0.1*cm))

        for i in range(0, len(valid_photos), 2):
            row_photos = valid_photos[i:i+2]
            row_data = []
            for photo, abs_p in row_photos:
                try:
                    img = RLImage(str(abs_p), width=8*cm, height=4.5*cm)
                    caption = f"{photo.caption or 'Site Photograph'}\n{str(photo.captured_at)[:16] if photo.captured_at else ''}"
                    cell = [img, Paragraph(to_pdf_html(caption), ParagraphStyle("Cap", fontSize=8, fontName='Helvetica', alignment=TA_CENTER))]
                except Exception as ex:
                    cell = [Paragraph(to_pdf_html(f"Failed to load photograph: {ex}"), normal)]
                row_data.append(cell)

            if len(row_data) == 1:
                row_data.append([Paragraph("", normal)])

            photo_table = Table([row_data], colWidths=[9*cm, 9*cm])
            photo_table.setStyle(TableStyle([
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("PADDING", (0, 0), (-1, -1), 3),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
            ]))
            story.append(photo_table)
            story.append(Spacer(1, 0.1*cm))

    # ── Approval Table ─────────────────────────────────────────
    if approvals:
        story.append(Spacer(1, 0.15*cm))
        story.append(Paragraph(to_pdf_html(L("Approval Details", "स्वीकृति विवरण"), bold=True), label))
        approval_headers = [
            Paragraph(to_pdf_html(L("Level", "स्तर"), bold=True), ParagraphStyle("H5", fontName='Helvetica-Bold', fontSize=9, textColor=colors.white)),
            Paragraph(to_pdf_html(L("Officer", "अधिकारी"), bold=True), ParagraphStyle("H5", fontName='Helvetica-Bold', fontSize=9, textColor=colors.white)),
            Paragraph(to_pdf_html(L("Status", "स्थिति"), bold=True), ParagraphStyle("H5", fontName='Helvetica-Bold', fontSize=9, textColor=colors.white)),
            Paragraph(to_pdf_html(L("Remarks", "टिप्पणी"), bold=True), ParagraphStyle("H5", fontName='Helvetica-Bold', fontSize=9, textColor=colors.white)),
            Paragraph(to_pdf_html(L("Date", "दिनांक"), bold=True), ParagraphStyle("H5", fontName='Helvetica-Bold', fontSize=9, textColor=colors.white))
        ]
        
        approval_data = [approval_headers]
        for a in approvals:
            app_name = a.approver.name_hindi or a.approver.name if a.approver else "N/A"
            desig = a.approver.designation or "Officer" if a.approver else ""
            act_labels = {"pending": L("Pending", "लंबित (Pending)"), "approved": L("Approved", "स्वीकृत (Approved)"), "rejected": L("Rejected", "अस्वीकृत (Rejected)"), "forwarded": L("Forwarded", "प्रेषित (Forwarded)")}
            action_val = get_val_str(a.action)
            action_eng = act_labels.get(action_val.lower(), action_val.upper())
            
            approval_data.append([
                Paragraph(to_pdf_html(a.level), normal),
                Paragraph(to_pdf_html(f"{app_name} ({desig})"), normal),
                Paragraph(to_pdf_html(action_eng), normal),
                Paragraph(to_pdf_html(a.remarks or "-"), normal),
                Paragraph(to_pdf_html(str(a.created_at)[:16]), normal),
            ])
        approval_table = Table(approval_data, colWidths=[2*cm, 4.5*cm, 2.5*cm, 6*cm, 3*cm])
        approval_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), DARK_BG),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("PADDING", (0, 0), (-1, -1), 3),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, LIGHT_BG]),
        ]))
        story.append(approval_table)

    # ── Signature ─────────────────────────────────────────────
    story.append(Spacer(1, 0.3*cm))
    witness_name = extract_witness_name(inspection.action_taken)
    if witness_name == "___________________":
        witness_name = extract_witness_name(inspection.description)
    engineer_designation = engineer.designation or "Junior Engineer" if engineer else "N/A"
    if engineer_designation.strip().lower() in ["super admin", "superadmin"]:
        engineer_designation = ""

    sig_data = [
        [Paragraph(to_pdf_html(L("Signature of Inspecting Officer", "निरीक्षण अधिकारी के हस्ताक्षर"), bold=True), normal), "", Paragraph(to_pdf_html(L("Signature of Witness / Representative", "गवाह / प्रतिनिधि के हस्ताक्षर"), bold=True), normal)],
        ["", "", ""],
        [Paragraph(to_pdf_html(L(f"Name: {engineer_name}", f"नाम: {engineer_name}")), normal), "",
         Paragraph(to_pdf_html(L(f"Name: {witness_name}", f"नाम: {witness_name}")), normal)],
        [Paragraph(to_pdf_html(L(f"Designation: {engineer_designation}", f"पदनाम: {engineer_designation}")), normal), "",
         Paragraph(to_pdf_html(""), normal)],
        [Paragraph(to_pdf_html(L(f"Date: {datetime.now().strftime('%d/%m/%Y')}", f"दिनांक: {datetime.now().strftime('%d/%m/%Y')}")), normal), "",
         Paragraph(to_pdf_html(L("Date: ___________________", "दिनांक: ___________________")), normal)],
    ]
    sig_table = Table(sig_data, colWidths=[8*cm, 3*cm, 8*cm])
    sig_table.setStyle(TableStyle([
        ("LINEABOVE", (0, 2), (0, 2), 1, colors.black),
        ("LINEABOVE", (2, 2), (2, 2), 1, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
    ]))
    story.append(sig_table)

    # ── Footer ────────────────────────────────────────────────
    story.append(Spacer(1, 0.15*cm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey))
    footer_text = f"Auto-generated by Gram Nirikshan Mobile App | {datetime.now().strftime('%d/%m/%Y %H:%M')} | Inspection ID: {inspection.inspection_id}"
    if has_hindi:
        footer_text = f"ग्राम निरीक्षण मोबाइल ऐप द्वारा स्वचालित रूप से उत्पन्न | {datetime.now().strftime('%d/%m/%Y %H:%M')} | निरीक्षण आईडी: {inspection.inspection_id}"
    story.append(Paragraph(
        to_pdf_html(footer_text),
        ParagraphStyle("Footer", fontSize=8, fontName='Helvetica', alignment=TA_CENTER, textColor=colors.grey)
    ))

    doc.build(story)


async def translate_fields_to_english(fields_dict: dict) -> dict:
    # Filter only fields that have Hindi characters (range 0x0900 to 0x097F)
    hindi_fields = {}
    for k, v in fields_dict.items():
        if v and isinstance(v, str) and any(0x0900 <= ord(c) <= 0x097F for c in v):
            hindi_fields[k] = v
            
    if not hindi_fields:
        return fields_dict
        
    try:
        from app.api.routes.ai import call_gemini
        import json
        
        prompt = f"""You are a professional translator. Translate the following fields from Hindi to standard, professional English. Keep technical terms, names, and addresses appropriate for a formal government report.
        
Provide your response strictly as a JSON object matching the input keys, with the translated values. Do not output any markdown formatting, no explanation, no backticks (e.g. do not wrap in ```json). Just output raw JSON.

Input JSON:
{json.dumps(hindi_fields, ensure_ascii=False)}"""

        translated_res = await call_gemini(prompt, language="en")
        if translated_res and not translated_res.startswith("AI Error:"):
            clean_res = translated_res.strip()
            if clean_res.startswith("```"):
                lines = clean_res.split("\n")
                if lines[0].startswith("```"):
                    lines = lines[1:]
                if lines[-1].strip() == "```":
                    lines = lines[:-1]
                clean_res = "\n".join(lines).strip()
            
            try:
                translated_dict = json.loads(clean_res)
                res_dict = dict(fields_dict)
                for k, v in translated_dict.items():
                    if k in res_dict:
                        res_dict[k] = v
                return res_dict
            except Exception as parse_err:
                import logging
                logging.getLogger(__name__).error(f"Failed to parse translation JSON: {parse_err}. Response was: {clean_res}")
    except Exception as e:
        import logging
        logging.getLogger(__name__).error(f"Bulk translation failed: {e}")
        
    return fields_dict


@router.post("/generate/{inspection_id}", response_model=MessageResponse)
async def generate_report(
    inspection_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Generate PDF report for an inspection."""
    result = await db.execute(select(Inspection).where(Inspection.id == inspection_id))
    inspection = result.scalar_one_or_none()
    if not inspection:
        raise HTTPException(status_code=404, detail="Inspection not found")

    result2 = await db.execute(select(Panchayat).where(Panchayat.id == inspection.panchayat_id))
    panchayat = result2.scalar_one_or_none()

    result3 = await db.execute(select(User).where(User.id == inspection.engineer_id))
    engineer = result3.scalar_one_or_none()

    result4 = await db.execute(select(Photo).where(Photo.inspection_id == inspection_id))
    photos = result4.scalars().all()

    result5 = await db.execute(
        select(Approval)
        .where(Approval.inspection_id == inspection_id)
        .options(selectinload(Approval.approver))
    )
    approvals = result5.scalars().all()

    file_name = f"Report_{inspection.inspection_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    output_path = str(REPORTS_DIR / file_name)

    orig_inspection_vals = {
        "title": inspection.title,
        "project_name": inspection.project_name,
        "observations": inspection.observations,
        "recommendations": inspection.recommendations,
        "action_taken": inspection.action_taken,
        "investigator_name": inspection.investigator_name,
        "checkin_address": inspection.checkin_address,
    }
    
    orig_panchayat_vals = {}
    if panchayat:
        orig_panchayat_vals["name"] = panchayat.name
        orig_panchayat_vals["name_hindi"] = panchayat.name_hindi
        orig_panchayat_vals["village"] = panchayat.village
        
    orig_engineer_vals = {}
    if engineer:
        orig_engineer_vals["name"] = engineer.name
        orig_engineer_vals["name_hindi"] = engineer.name_hindi
        orig_engineer_vals["designation"] = engineer.designation

    orig_photo_captions = {p.id: p.caption for p in photos if p.id}
    orig_approval_remarks = {a.id: a.remarks for a in approvals if a.id}

    try:
        # Collect all fields to translate in a single bulk API call
        fields_to_translate = {
            "title": inspection.title,
            "project_name": inspection.project_name,
            "observations": inspection.observations,
            "recommendations": inspection.recommendations,
            "action_taken": inspection.action_taken,
            "investigator_name": inspection.investigator_name,
            "checkin_address": inspection.checkin_address,
        }
        if panchayat:
            fields_to_translate["panchayat_name"] = panchayat.name
            fields_to_translate["panchayat_name_hindi"] = panchayat.name_hindi
            fields_to_translate["panchayat_village"] = panchayat.village
        if engineer:
            fields_to_translate["engineer_name"] = engineer.name
            fields_to_translate["engineer_name_hindi"] = engineer.name_hindi
            fields_to_translate["engineer_designation"] = engineer.designation

        for p in photos:
            if p.caption:
                fields_to_translate[f"photo_{p.id}_caption"] = p.caption

        for a in approvals:
            if a.remarks:
                fields_to_translate[f"approval_{a.id}_remarks"] = a.remarks

        # Determine language dynamically based on input fields containing Hindi characters
        has_hindi = any(v and isinstance(v, str) and any(0x0900 <= ord(c) <= 0x097F for c in v) 
                        for v in [inspection.title, inspection.project_name, inspection.observations, inspection.recommendations])

        if has_hindi:
            # Bypass translation if Hindi is detected
            translated_fields = fields_to_translate
        else:
            translated_fields = await translate_fields_to_english(fields_to_translate)

        # Assign translated values back to objects
        inspection.title = translated_fields.get("title")
        inspection.project_name = translated_fields.get("project_name")
        inspection.observations = translated_fields.get("observations")
        inspection.recommendations = translated_fields.get("recommendations")
        inspection.action_taken = translated_fields.get("action_taken")
        inspection.investigator_name = translated_fields.get("investigator_name")
        inspection.checkin_address = translated_fields.get("checkin_address")

        if panchayat:
            panchayat.name = translated_fields.get("panchayat_name")
            panchayat.name_hindi = translated_fields.get("panchayat_name_hindi")
            panchayat.village = translated_fields.get("panchayat_village")

        if engineer:
            engineer.name = translated_fields.get("engineer_name")
            engineer.name_hindi = translated_fields.get("engineer_name_hindi")
            engineer.designation = translated_fields.get("engineer_designation")

        for p in photos:
            key = f"photo_{p.id}_caption"
            if key in translated_fields:
                p.caption = translated_fields[key]

        for a in approvals:
            key = f"approval_{a.id}_remarks"
            if key in translated_fields:
                a.remarks = translated_fields[key]

        # Check if AI report draft is missing, generate it dynamically using Gemini
        if not inspection.ai_report_draft:
            try:
                from app.api.routes.ai import call_gemini
                
                if has_hindi:
                    prompt = f"""Draft a highly formal and professional Gram Panchayat inspection report (Inspection Memo) in Hindi according to the standards of the Rural Development Department (Gram Panchayat Department).

Inspection Details:
- Inspection ID: {inspection.inspection_id}
- Title: {inspection.title}
- Gram Panchayat: {panchayat.name_hindi or panchayat.name if panchayat else 'N/A'} (District: {inspection.district or (panchayat.district if panchayat else 'N/A')}, Block: {inspection.block or (panchayat.block if panchayat else 'N/A')})
- Inspector/Engineer: {inspection.investigator_name or (engineer.name_hindi or engineer.name if engineer else 'N/A')} (Designation: {engineer.designation or 'Junior Engineer' if engineer else 'N/A'})
- Project/Work Name: {inspection.project_name or 'N/A'} (Work Code: {inspection.project_code or 'N/A'})
- Inspection Type: {inspection.inspection_type or 'General'}

Observations / Notes:
{inspection.observations or 'निरीक्षण किया गया।'}

Corrective Recommendations:
{inspection.recommendations or 'उचित सुधारात्मक कार्रवाई की जानी चाहिए।'}

Draft the full Hindi report under the following sections:
1. **कार्य का विवरण और मुख्य निष्कर्ष (क्या अच्छा था)**: निरीक्षण का विवरण, कार्य की प्रगति और सकारात्मक निष्कर्ष।
2. **कमियां / पहचानी गई समस्याएं (क्या कमी थी)**: निरीक्षण के दौरान पाई गई तकनीकी, गुणवत्ता संबंधी या प्रशासनिक कमियां।
3. **सुधारात्मक कार्रवाई / सिफारिशें (कैसे समाधान किया जाए)**: पहचानी गई कमियों को दूर करने के लिए आवश्यक कार्रवाई और सिफारिशें।
4. **निष्कर्ष**: कार्य की गुणवत्ता पर अंतिम टिप्पणी और आगे के कदम।

Ensure the report is professional, grammatically correct, and written in clear technical standard Hindi suitable for senior administration."""
                    ai_draft = await call_gemini(prompt, language="hi")
                else:
                    prompt = f"""Draft a highly formal and professional Gram Panchayat inspection report (Inspection Memo) in English according to the standards of the Rural Development Department.

Inspection Details:
- Inspection ID: {inspection.inspection_id}
- Title: {inspection.title}
- Gram Panchayat: {panchayat.name or panchayat.name_hindi if panchayat else 'N/A'} (District: {inspection.district or (panchayat.district if panchayat else 'N/A')}, Block: {inspection.block or (panchayat.block if panchayat else 'N/A')})
- Inspector/Engineer: {inspection.investigator_name or (engineer.name or engineer.name_hindi if engineer else 'N/A')} (Designation: {engineer.designation or 'Junior Engineer' if engineer else 'N/A'})
- Project/Work Name: {inspection.project_name or 'N/A'} (Work Code: {inspection.project_code or 'N/A'})
- Inspection Type: {inspection.inspection_type or 'General'}

Observations / Notes:
{inspection.observations or 'Site inspection conducted.'}

Corrective Recommendations:
{inspection.recommendations or 'Appropriate corrective measures should be taken.'}

Draft the full English report under the following sections:
1. **Work Description & Key Findings (What was good)**: Details of the site inspection, work progress, and positive findings.
2. **Deficiencies / Issues Identified (What was lacking)**: Technical, quality-related, or administrative deficiencies observed during the inspection.
3. **Corrective Actions / Recommendations (What can be resolved)**: Necessary corrective actions and recommendations to resolve the identified deficiencies.
4. **Conclusion**: Final remarks on work quality and next steps.

Ensure the report is professional, grammatically correct, and written in clear technical English suitable for senior administration."""
                    ai_draft = await call_gemini(prompt, language="en")
                
                if ai_draft and not ai_draft.startswith("AI Error:"):
                    inspection.ai_report_draft = ai_draft
                    await db.flush()
            except Exception as e:
                import logging
                logging.getLogger(__name__).error(f"Failed to auto-generate Gemini report draft: {e}")

        try:
            build_pdf_report(inspection, panchayat, engineer, list(photos), list(approvals), output_path)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")

        # Save PDF report record
        file_size_kb = Path(output_path).stat().st_size // 1024
        report = Report(
            inspection_id=inspection_id,
            generated_by=current_user.id,
            file_path=output_path,
            file_name=file_name,
            file_size_kb=file_size_kb,
            report_format="pdf",
        )
        db.add(report)

        # DOCX Generation
        docx_file_name = f"Report_{inspection.inspection_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        docx_output_path = str(REPORTS_DIR / docx_file_name)
        try:
            build_docx_report(inspection, panchayat, engineer, list(photos), list(approvals), docx_output_path)
            docx_file_size_kb = Path(docx_output_path).stat().st_size // 1024
            docx_report = Report(
                inspection_id=inspection_id,
                generated_by=current_user.id,
                file_path=docx_output_path,
                file_name=docx_file_name,
                file_size_kb=docx_file_size_kb,
                report_format="docx",
            )
            db.add(docx_report)
        except Exception as docx_err:
            import logging
            logging.getLogger(__name__).error(f"DOCX generation failed: {docx_err}")

    finally:
        # Restore original values
        for k, v in orig_inspection_vals.items():
            setattr(inspection, k, v)
        if panchayat:
            for k, v in orig_panchayat_vals.items():
                setattr(panchayat, k, v)
        if engineer:
            for k, v in orig_engineer_vals.items():
                setattr(engineer, k, v)
        for p in photos:
            if p.id in orig_photo_captions:
                p.caption = orig_photo_captions[p.id]
        for a in approvals:
            if a.id in orig_approval_remarks:
                a.remarks = orig_approval_remarks[a.id]

    return MessageResponse(
        message="Reports generated successfully",
        success=True,
        data={"file_name": file_name, "size_kb": file_size_kb},
    )


@router.get("/download/{inspection_id}")
async def download_report(
    inspection_id: str,
    format: str = "pdf",
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Download latest report PDF or DOCX for an inspection."""
    result = await db.execute(
        select(Report).where(Report.inspection_id == inspection_id)
        .where(Report.report_format == format)
        .order_by(Report.created_at.desc())
    )
    report = result.scalars().first()
    if not report or not Path(report.file_path).exists():
        raise HTTPException(status_code=404, detail=f"Report in {format} format not found. Generate it first.")

    media_type = "application/pdf" if format == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return FileResponse(
        report.file_path,
        media_type=media_type,
        filename=report.file_name,
    )

