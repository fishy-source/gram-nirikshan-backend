"""
PDF Report Generator using ReportLab.
Generates department-format inspection reports with photos, GPS data, and approval trail.
"""
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, inch
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    Image as RLImage, HRFlowable, KeepTogether
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from pathlib import Path
from datetime import datetime, timezone
import uuid, os, shutil

try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    docx = None

from app.db.database import get_db
from app.models.models import Inspection, Panchayat, User, Photo, Approval, Report
from app.schemas.schemas import MessageResponse
from app.core.dependencies import get_current_user
from app.core.config import settings

router = APIRouter(prefix="/reports", tags=["Reports"])

REPORTS_DIR = Path(settings.UPLOAD_DIR) / "reports"
REPORTS_DIR.mkdir(parents=True, exist_ok=True)

# Color scheme
PRIMARY = colors.HexColor("#1a5276")
SECONDARY = colors.HexColor("#2e86c1")
ACCENT = colors.HexColor("#f39c12")
LIGHT_BG = colors.HexColor("#eaf2fb")
DARK_BG = colors.HexColor("#154360")


def find_project_root() -> Path:
    current = Path(__file__).resolve().parent
    for _ in range(5):
        if (current / "flutter_app").exists():
            return current
        current = current.parent
    return Path(__file__).parents[3]


def get_absolute_path(rel_path: str) -> Path:
    """Find absolute path of a file, searching workspace directories."""
    path = Path(rel_path)
    if path.exists():
        return path
    root = find_project_root()
    path2 = root / rel_path
    if path2.exists():
        return path2
    path3 = root / "backend" / rel_path
    if path3.exists():
        return path3
    return path


# Register Poppins Hindi Unicode font
pdf_font_name = 'Helvetica'
pdf_font_bold_name = 'Helvetica-Bold'

try:
    root = find_project_root()
    flutter_font_dir = root / "flutter_app" / "assets" / "fonts"
    backend_font_dir = root / "assets" / "fonts"
    backend_font_dir.mkdir(parents=True, exist_ok=True)

    for font_name in ["Poppins-Regular.ttf", "Poppins-Bold.ttf"]:
        src = flutter_font_dir / font_name
        dst = backend_font_dir / font_name
        if src.exists() and not dst.exists():
            try:
                shutil.copy(src, dst)
            except Exception as ce:
                print(f"Failed to copy font {font_name}: {ce}")

    # Verify and register from candidate paths
    possible_paths = [
        (backend_font_dir / "Poppins-Regular.ttf", backend_font_dir / "Poppins-Bold.ttf"),
        (flutter_font_dir / "Poppins-Regular.ttf", flutter_font_dir / "Poppins-Bold.ttf"),
        (Path("/app/flutter_app/assets/fonts/Poppins-Regular.ttf"), Path("/app/flutter_app/assets/fonts/Poppins-Bold.ttf")),
        (Path("/app/assets/fonts/Poppins-Regular.ttf"), Path("/app/assets/fonts/Poppins-Bold.ttf")),
    ]

    font_path_regular = None
    font_path_bold = None

    for reg, bld in possible_paths:
        if reg.exists() and bld.exists():
            font_path_regular = reg
            font_path_bold = bld
            break

    if font_path_regular and font_path_bold:
        pdfmetrics.registerFont(TTFont('Poppins', str(font_path_regular)))
        pdfmetrics.registerFont(TTFont('Poppins-Bold', str(font_path_bold)))
        pdf_font_name = 'Poppins'
        pdf_font_bold_name = 'Poppins-Bold'
except Exception as e:
    print(f"Error registering Poppins fonts: {e}")
    pdf_font_name = 'Helvetica'
    pdf_font_bold_name = 'Helvetica-Bold'


def extract_witness_name(description: str) -> str:
    if not description:
        return "___________________"
    import re
    match = re.search(r"(?:गवाह का नाम|गवाह\s*नाम|गवाह|witness name|witness)\s*[:：\-‐—\s]\s*([^\n\r।\.]+)", description, re.IGNORECASE)
    if match:
        name = match.group(1).strip()
        name = name.split(',')[0].split(';')[0].strip()
        name = re.sub(r"^(?:श्री|shri|mr\.|mrs\.)\s+", "", name, flags=re.IGNORECASE)
        stop_words = ["उपस्थित", "मौजूद", "थे", "था", "रहा", "रही", "गया", "गई", "है", "हैं", "present", "here", "was", "is", "were"]
        for stop_word in stop_words:
            name = re.sub(rf"(?:\s+|^){stop_word}(?:\s+|$)", " ", name, flags=re.IGNORECASE).strip()
        if name:
            return name
    return "___________________"


def build_docx_report(inspection, panchayat, engineer, photos, approvals, output_path: str):
    """Build a formal Microsoft Word report."""
    if not docx:
        raise RuntimeError("python-docx library is not installed on this server.")
        
    doc = docx.Document()
    
    # Set margins to 0.75 inches for page-like structure
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    # Title / Header
    dept_name = settings.DEPARTMENT_NAME or "ग्राम विकास विभाग, उत्तर प्रदेश"
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(dept_name)
    run_title.font.size = Pt(16)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(0x1a, 0x52, 0x76) # PRIMARY
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("निरीक्षण रिपोर्ट (INSPECTION REPORT)\n")
    run_sub.font.size = Pt(11)
    run_sub.bold = True
    run_sub.font.color.rgb = RGBColor(0x2e, 0x86, 0xc1) # SECONDARY
    
    run_id = p_sub.add_run(f"रिपोर्ट संख्या (Report ID): {inspection.inspection_id}")
    run_id.font.size = Pt(10)
    run_id.italic = True
    
    # Border line
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run("─" * 65)
    run_line.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)
    
    # Basic Information
    engineer_name = inspection.investigator_name or (engineer.name_hindi or engineer.name if engineer else "N/A")
    dist = inspection.district or (panchayat.district if panchayat else "N/A")
    blk = inspection.block or (panchayat.block if panchayat else "N/A")
    
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    
    def set_cell(cell, text, bold=False, text_color=None):
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run(text)
        run.font.size = Pt(10)
        if bold:
            run.bold = True
        if text_color:
            run.font.color.rgb = text_color
            
    # Row contents helper
    info_pairs = [
        (f"निरीक्षण संख्या (ID): {inspection.inspection_id}", f"स्थिति (Status): {inspection.status.value.upper()}"),
        (f"जांचकर्ता (Inspector): {engineer_name}", f"पद (Designation): {engineer.designation or 'अवर अभियंता' if engineer else 'N/A'}"),
        (f"ग्राम पंचायत: {panchayat.name_hindi or panchayat.name if panchayat else 'N/A'}", f"जनपद (District): {dist}"),
        (f"विकास खंड (Block): {blk}", f"ग्राम (Village): {panchayat.village or 'N/A' if panchayat else 'N/A'}"),
        (f"निरीक्षण तिथि: {str(inspection.inspection_date)[:10] if inspection.inspection_date else 'N/A'}", f"रिपोर्ट तिथि: {datetime.now().strftime('%d/%m/%Y')}"),
        (f"परियोजना का नाम: {inspection.project_name or 'N/A'}", f"कार्य कोड (Code): {inspection.project_code or 'N/A'}"),
    ]
    
    for idx, (col1, col2) in enumerate(info_pairs):
        row = table.rows[idx]
        set_cell(row.cells[0], col1)
        set_cell(row.cells[1], col2)
        
    doc.add_paragraph() # Spacer
    
    # GPS Details
    if inspection.checkin_latitude:
        p_gps_lbl = doc.add_paragraph()
        run_gps_lbl = p_gps_lbl.add_run("जीपीएस विवरण (GPS Check-in/Check-out Details)")
        run_gps_lbl.bold = True
        run_gps_lbl.font.size = Pt(11)
        run_gps_lbl.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)
        
        gps_table = doc.add_table(rows=3, cols=2)
        gps_table.style = 'Table Grid'
        
        gps_pairs = [
            (f"चेक-इन समय: {str(inspection.checkin_time)[:16] if inspection.checkin_time else 'N/A'}", f"चेक-इन जीपीएस: {inspection.checkin_latitude:.6f}, {inspection.checkin_longitude:.6f}"),
            (f"चेक-आउट समय: {str(inspection.checkout_time)[:16] if inspection.checkout_time else 'N/A'}", f"चेक-आउट जीपीएस: {f'{inspection.checkout_latitude:.6f}, {inspection.checkout_longitude:.6f}' if inspection.checkout_latitude else 'N/A'}"),
            (f"दूरी (Distance): {f'{inspection.distance_covered_km:.2f} किमी' if inspection.distance_covered_km else 'N/A'}", f"चेक-इन स्थान: {inspection.checkin_address or 'N/A'}"),
        ]
        
        for idx, (col1, col2) in enumerate(gps_pairs):
            row = gps_table.rows[idx]
            set_cell(row.cells[0], col1)
            set_cell(row.cells[1], col2)
            
        doc.add_paragraph() # Spacer
        
    # Map Attachment
    if inspection.map_image_path:
        map_file_path = get_absolute_path(inspection.map_image_path)
        if map_file_path.exists():
            p_map = doc.add_paragraph()
            run_map = p_map.add_run("निरीक्षण स्थान मानचित्र (Inspection Location Map)")
            run_map.bold = True
            run_map.font.size = Pt(11)
            run_map.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)
            
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                p_img.add_run().add_picture(str(map_file_path), width=Inches(5.5))
            except Exception as e:
                p_img.add_run(f"मानचित्र छवि लोड करने में विफल: {e}")
            doc.add_paragraph()
            
    # Observations & Recommendations
    for section_title, content in [
        ("मुख्य अवलोकन / कमियां (Observations)", inspection.observations),
        ("सुझाव व संस्तुतियां (Recommendations)", inspection.recommendations),
        ("की गई कार्रवाई (Action Taken)", inspection.action_taken),
        ("विभागीय मसौदा रिपोर्ट (AI Suggested Report)", inspection.ai_report_draft),
    ]:
        if content:
            p_lbl = doc.add_paragraph()
            run_lbl = p_lbl.add_run(section_title)
            run_lbl.bold = True
            run_lbl.font.size = Pt(11)
            run_lbl.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)
            
            p_val = doc.add_paragraph()
            p_val.paragraph_format.space_after = Pt(6)
            run_val = p_val.add_run(content)
            run_val.font.size = Pt(10)
            doc.add_paragraph() # Spacer
            
    # Photos
    valid_photos = []
    for p in photos:
        if p.file_path:
            abs_p = get_absolute_path(p.file_path)
            if abs_p.exists():
                valid_photos.append((p, abs_p))
                
    if valid_photos:
        p_photo_lbl = doc.add_paragraph()
        run_photo_lbl = p_photo_lbl.add_run("निरीक्षण स्थल के छायाचित्र (Inspection Photographs)")
        run_photo_lbl.bold = True
        run_photo_lbl.font.size = Pt(11)
        run_photo_lbl.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)
        
        # Grid layout for photos: standard tables can embed pictures
        photo_table = doc.add_table(rows=(len(valid_photos) + 1) // 2, cols=2)
        photo_table.style = 'Table Grid'
        
        for idx, (photo, abs_p) in enumerate(valid_photos):
            row_idx = idx // 2
            col_idx = idx % 2
            cell = photo_table.rows[row_idx].cells[col_idx]
            cell.text = ""
            p_cell = cell.paragraphs[0]
            p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                p_cell.add_run().add_picture(str(abs_p), width=Inches(2.5))
                caption_text = f"\n{photo.caption or 'स्थल छायाचित्र'}\n{str(photo.captured_at)[:16] if photo.captured_at else ''}"
                run_cap = p_cell.add_run(caption_text)
                run_cap.font.size = Pt(8)
                run_cap.italic = True
            except Exception as e:
                p_cell.add_run(f"छायाचित्र लोड करने में विफल: {e}")
        doc.add_paragraph()
        
    # Approval Trail
    if approvals:
        p_app_lbl = doc.add_paragraph()
        run_app_lbl = p_app_lbl.add_run("कार्यप्रवाह एवं अनुमोदन इतिहास (Approval Trail)")
        run_app_lbl.bold = True
        run_app_lbl.font.size = Pt(11)
        run_app_lbl.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)
        
        app_table = doc.add_table(rows=1 + len(approvals), cols=5)
        app_table.style = 'Table Grid'
        
        headers = ["स्तर (Level)", "अधिकारी (Approver)", "कार्रवाई (Action)", "टिप्पणी (Remarks)", "दिनांक (Date)"]
        for col_idx, text in enumerate(headers):
            set_cell(app_table.rows[0].cells[col_idx], text, bold=True, text_color=RGBColor(0xff, 0xff, 0xff))
            
        for row_idx, a in enumerate(approvals):
            app_name = a.approver.name_hindi or a.approver.name if a.approver else "N/A"
            desig = a.approver.designation or "अधिकारी" if a.approver else ""
            act_labels = {"pending": "लंबित", "approved": "स्वीकृत", "rejected": "अस्वीकृत", "forwarded": "अग्रेषित"}
            action_hindi = act_labels.get(a.action.value.lower(), a.action.value.upper())
            
            row = app_table.rows[row_idx + 1]
            set_cell(row.cells[0], a.level)
            set_cell(row.cells[1], f"{app_name} ({desig})")
            set_cell(row.cells[2], action_hindi)
            set_cell(row.cells[3], a.remarks or "-")
            set_cell(row.cells[4], str(a.created_at)[:16])
            
        doc.add_paragraph()
        
    # Signatures
    doc.add_paragraph("\n\n")
    witness_name = extract_witness_name(inspection.description)
    engineer_designation = engineer.designation or "अवर अभियंता" if engineer else "N/A"
    
    p_sig = doc.add_paragraph()
    p_sig.paragraph_format.line_spacing = 1.3
    
    run_sig_h = p_sig.add_run("जांचकर्ता अधिकारी के हस्ताक्षर" + " " * 32 + "गवाह के हस्ताक्षर\n")
    run_sig_h.bold = True
    run_sig_h.font.size = Pt(10)
    
    engineer_part = f"नाम (Name): {engineer_name}"
    witness_part = f"नाम (Name): {witness_name}"
    spacing = max(5, 50 - len(engineer_part))
    p_sig.add_run(engineer_part + " " * spacing + witness_part + "\n")
    
    desig_part = f"पदनाम (Designation): {engineer_designation}"
    addr_part = "पता/मोबाइल: ___________________"
    spacing_desig = max(5, 50 - len(desig_part))
    p_sig.add_run(desig_part + " " * spacing_desig + addr_part + "\n")
    
    date_part = f"दिनांक (Date): {datetime.now().strftime('%d/%m/%Y')}"
    date_witness = "दिनांक (Date): ___________________"
    spacing_date = max(5, 50 - len(date_part))
    p_sig.add_run(date_part + " " * spacing_date + date_witness + "\n")
    
    for r in p_sig.runs[1:]:
        r.font.size = Pt(10)
        
    # Footer
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_foot = p_foot.add_run(f"\n──────────────────────────────────────────────────\nग्राम निरीक्षण मोबाइल ऐप द्वारा स्वचालित जनरेटेड | {datetime.now().strftime('%d/%m/%Y %H:%M')} | निरीक्षण ID: {inspection.inspection_id}")
    run_foot.font.size = Pt(8)
    run_foot.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    
    doc.save(output_path)


def build_pdf_report(inspection, panchayat, engineer, photos, approvals, output_path: str):
    """Build a full PDF inspection report."""
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2.5*cm, bottomMargin=2*cm,
    )

    styles = getSampleStyleSheet()
    story = []

    # ── Header ────────────────────────────────────────────────
    title_style = ParagraphStyle(
        "Title", fontSize=18, textColor=PRIMARY,
        alignment=TA_CENTER, fontName=pdf_font_bold_name, spaceAfter=4
    )
    subtitle_style = ParagraphStyle(
        "Subtitle", fontSize=11, textColor=SECONDARY,
        alignment=TA_CENTER, fontName=pdf_font_name, spaceAfter=2
    )
    normal = ParagraphStyle("Normal2", fontSize=10, fontName=pdf_font_name, spaceAfter=4)
    label = ParagraphStyle("Label", fontSize=10, fontName=pdf_font_bold_name, textColor=PRIMARY)

    dept_name = settings.DEPARTMENT_NAME or "ग्राम विकास विभाग, उत्तर प्रदेश"
    story.append(Paragraph(dept_name, title_style))
    story.append(Paragraph("निरीक्षण रिपोर्ट (INSPECTION REPORT)", subtitle_style))
    story.append(Paragraph(f"रिपोर्ट संख्या (Report ID): {inspection.inspection_id}", subtitle_style))
    story.append(HRFlowable(width="100%", thickness=2, color=PRIMARY, spaceAfter=12))

    # ── Basic Information ──────────────────────────────────────
    engineer_name = inspection.investigator_name or (engineer.name_hindi or engineer.name if engineer else "N/A")
    dist = inspection.district or (panchayat.district if panchayat else "N/A")
    blk = inspection.block or (panchayat.block if panchayat else "N/A")

    info_data = [
        [Paragraph("निरीक्षण संख्या (ID)", normal), Paragraph(inspection.inspection_id, normal), 
         Paragraph("स्थिति (Status)", normal), Paragraph(inspection.status.value.upper(), normal)],
        [Paragraph("जांचकर्ता (Inspector)", normal), Paragraph(engineer_name, normal), 
         Paragraph("पद (Designation)", normal), Paragraph(engineer.designation or "अवर अभियंता" if engineer else "N/A", normal)],
        [Paragraph("ग्राम पंचायत", normal), Paragraph(panchayat.name_hindi or panchayat.name if panchayat else "N/A", normal), 
         Paragraph("जनपद (District)", normal), Paragraph(dist, normal)],
        [Paragraph("विकास खंड (Block)", normal), Paragraph(blk, normal), 
         Paragraph("ग्राम (Village)", normal), Paragraph(panchayat.village or "N/A" if panchayat else "N/A", normal)],
        [Paragraph("निरीक्षण तिथि", normal), Paragraph(str(inspection.inspection_date)[:10] if inspection.inspection_date else "N/A", normal),
         Paragraph("रिपोर्ट तिथि", normal), Paragraph(datetime.now(timezone.utc).strftime("%d/%m/%Y"), normal)],
        [Paragraph("परियोजना का नाम", normal), Paragraph(inspection.project_name or "N/A", normal), 
         Paragraph("कार्य कोड (Code)", normal), Paragraph(inspection.project_code or "N/A", normal)],
    ]

    info_table = Table(info_data, colWidths=[4*cm, 5.5*cm, 4*cm, 5.5*cm])
    info_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), LIGHT_BG),
        ("BACKGROUND", (2, 0), (2, -1), LIGHT_BG),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("PADDING", (0, 0), (-1, -1), 6),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 0.3*cm))

    # ── GPS Information ────────────────────────────────────────
    if inspection.checkin_latitude:
        story.append(Paragraph("जीपीएस विवरण (GPS Check-in/Check-out Details)", label))
        gps_data = [
            [Paragraph("चेक-इन समय", normal), Paragraph(str(inspection.checkin_time)[:16] if inspection.checkin_time else "N/A", normal),
             Paragraph("चेक-इन जीपीएस", normal), Paragraph(f"{inspection.checkin_latitude:.6f}, {inspection.checkin_longitude:.6f}", normal)],
            [Paragraph("चेक-आउट समय", normal), Paragraph(str(inspection.checkout_time)[:16] if inspection.checkout_time else "N/A", normal),
             Paragraph("चेक-आउट जीपीएस", normal), Paragraph(f"{inspection.checkout_latitude:.6f}, {inspection.checkout_longitude:.6f}" if inspection.checkout_latitude else "N/A", normal)],
            [Paragraph("दूरी (Distance)", normal), Paragraph(f"{inspection.distance_covered_km:.2f} किमी" if inspection.distance_covered_km else "N/A", normal),
             Paragraph("चेक-इन स्थान", normal), Paragraph(inspection.checkin_address or "N/A", normal)],
        ]
        gps_table = Table(gps_data, colWidths=[4*cm, 5.5*cm, 4*cm, 5.5*cm])
        gps_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), LIGHT_BG),
            ("BACKGROUND", (2, 0), (2, -1), LIGHT_BG),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("PADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(gps_table)
        story.append(Spacer(1, 0.3*cm))

    # ── Map Attachment ─────────────────────────────────────────
    if inspection.map_image_path:
        map_file_path = get_absolute_path(inspection.map_image_path)
        if map_file_path.exists():
            story.append(Paragraph("निरीक्षण स्थान मानचित्र (Inspection Location Map)", label))
            story.append(Spacer(1, 0.2*cm))
            try:
                img = RLImage(str(map_file_path), width=16*cm, height=9*cm)
                story.append(KeepTogether([img, Spacer(1, 0.3*cm)]))
            except Exception as e:
                story.append(Paragraph(f"मानचित्र छवि लोड करने में विफल: {str(e)}", normal))

    # ── Observations ──────────────────────────────────────────
    for section_title, content in [
        ("मुख्य अवलोकन / कमियां (Observations)", inspection.observations),
        ("सुझाव व संस्तुतियां (Recommendations)", inspection.recommendations),
        ("की गई कार्रवाई (Action Taken)", inspection.action_taken),
        ("विभागीय मसौदा रिपोर्ट (AI Suggested Report)", inspection.ai_report_draft),
    ]:
        if content:
            story.append(Paragraph(section_title, label))
            story.append(Paragraph(content, normal))
            story.append(Spacer(1, 0.3*cm))

    # ── Photos ────────────────────────────────────────────────
    valid_photos = []
    for p in photos:
        if p.file_path:
            abs_p = get_absolute_path(p.file_path)
            if abs_p.exists():
                valid_photos.append((p, abs_p))

    if valid_photos:
        story.append(HRFlowable(width="100%", thickness=1, color=SECONDARY, spaceBefore=8, spaceAfter=8))
        story.append(Paragraph("निरीक्षण स्थल के छायाचित्र (Inspection Photographs)", label))
        story.append(Spacer(1, 0.2*cm))

        # 2 photos per row
        for i in range(0, len(valid_photos), 2):
            row_photos = valid_photos[i:i+2]
            row_data = []
            for photo, abs_p in row_photos:
                try:
                    img = RLImage(str(abs_p), width=8*cm, height=6*cm)
                    caption = f"{photo.caption or 'स्थल छायाचित्र'}\n{str(photo.captured_at)[:16] if photo.captured_at else ''}"
                    cell = [img, Paragraph(caption, ParagraphStyle("Cap", fontSize=8, fontName=pdf_font_name, alignment=TA_CENTER))]
                except Exception as ex:
                    cell = [Paragraph(f"छायाचित्र लोड करने में विफल: {ex}", normal)]
                row_data.append(cell)

            if len(row_data) == 1:
                row_data.append([""])  # Fill empty cell

            photo_table = Table([row_data], colWidths=[9*cm, 9*cm])
            photo_table.setStyle(TableStyle([
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("PADDING", (0, 0), (-1, -1), 5),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
            ]))
            story.append(photo_table)
            story.append(Spacer(1, 0.2*cm))

    # ── Approval Trail ────────────────────────────────────────
    if approvals:
        story.append(HRFlowable(width="100%", thickness=1, color=SECONDARY, spaceBefore=8, spaceAfter=8))
        story.append(Paragraph("कार्यप्रवाह एवं अनुमोदन इतिहास (Approval Trail)", label))
        story.append(Spacer(1, 0.2*cm))
        
        approval_headers = [
            Paragraph("स्तर (Level)", ParagraphStyle("H1", fontName=pdf_font_bold_name, fontSize=9, textColor=colors.white)),
            Paragraph("अधिकारी (Approver)", ParagraphStyle("H2", fontName=pdf_font_bold_name, fontSize=9, textColor=colors.white)),
            Paragraph("कार्रवाई (Action)", ParagraphStyle("H3", fontName=pdf_font_bold_name, fontSize=9, textColor=colors.white)),
            Paragraph("टिप्पणी (Remarks)", ParagraphStyle("H4", fontName=pdf_font_bold_name, fontSize=9, textColor=colors.white)),
            Paragraph("दिनांक (Date)", ParagraphStyle("H5", fontName=pdf_font_bold_name, fontSize=9, textColor=colors.white))
        ]
        
        approval_data = [approval_headers]
        for a in approvals:
            app_name = a.approver.name_hindi or a.approver.name if a.approver else "N/A"
            desig = a.approver.designation or "अधिकारी" if a.approver else ""
            act_labels = {"pending": "लंबित", "approved": "स्वीकृत", "rejected": "अस्वीकृत", "forwarded": "अग्रेषित"}
            action_hindi = act_labels.get(a.action.value.lower(), a.action.value.upper())
            
            approval_data.append([
                Paragraph(a.level, normal),
                Paragraph(f"{app_name} ({desig})", normal),
                Paragraph(action_hindi, normal),
                Paragraph(a.remarks or "-", normal),
                Paragraph(str(a.created_at)[:16], normal),
            ])
        approval_table = Table(approval_data, colWidths=[2*cm, 4.5*cm, 2.5*cm, 6*cm, 3*cm])
        approval_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), DARK_BG),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("PADDING", (0, 0), (-1, -1), 5),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, LIGHT_BG]),
        ]))
        story.append(approval_table)

    # ── Signature ─────────────────────────────────────────────
    story.append(Spacer(1, 1*cm))
    witness_name = extract_witness_name(inspection.description)
    engineer_designation = engineer.designation or "अवर अभियंता" if engineer else "N/A"

    sig_data = [
        [Paragraph("जांचकर्ता अधिकारी के हस्ताक्षर", normal), "", Paragraph("गवाह के हस्ताक्षर", normal)],
        ["", "", ""],
        [Paragraph(f"नाम (Name): {engineer_name}", normal), "",
         Paragraph(f"नाम (Name): {witness_name}", normal)],
        [Paragraph(f"पदनाम (Designation): {engineer_designation}", normal), "",
         Paragraph("पता/मोबाइल: ___________________", normal)],
        [Paragraph(f"दिनांक (Date): {datetime.now().strftime('%d/%m/%Y')}", normal), "",
         Paragraph("दिनांक (Date): ___________________", normal)],
    ]
    sig_table = Table(sig_data, colWidths=[8*cm, 3*cm, 8*cm])
    sig_table.setStyle(TableStyle([
        ("LINEABOVE", (0, 2), (0, 2), 1, colors.black),
        ("LINEABOVE", (2, 2), (2, 2), 1, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
    ]))
    story.append(sig_table)

    # ── Footer ────────────────────────────────────────────────
    story.append(Spacer(1, 0.5*cm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey))
    story.append(Paragraph(
        f"ग्राम निरीक्षण मोबाइल ऐप द्वारा स्वचालित जनरेटेड | {datetime.now().strftime('%d/%m/%Y %H:%M')} | निरीक्षण ID: {inspection.inspection_id}",
        ParagraphStyle("Footer", fontSize=8, fontName=pdf_font_name, alignment=TA_CENTER, textColor=colors.grey)
    ))

    doc.build(story)


@router.post("/generate/{inspection_id}", response_model=MessageResponse)
async def generate_report(
    inspection_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Generate PDF report for an inspection."""
    result = await db.execute(select(Inspection).where(Inspection.id == inspection_id))
    inspection = result.scalar_one_or_none()
    if not inspection:
        raise HTTPException(status_code=404, detail="Inspection not found")

    result2 = await db.execute(select(Panchayat).where(Panchayat.id == inspection.panchayat_id))
    panchayat = result2.scalar_one_or_none()

    result3 = await db.execute(select(User).where(User.id == inspection.engineer_id))
    engineer = result3.scalar_one_or_none()

    result4 = await db.execute(select(Photo).where(Photo.inspection_id == inspection_id))
    photos = result4.scalars().all()

    result5 = await db.execute(
        select(Approval)
        .where(Approval.inspection_id == inspection_id)
        .options(selectinload(Approval.approver))
    )
    approvals = result5.scalars().all()

    file_name = f"Report_{inspection.inspection_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    output_path = str(REPORTS_DIR / file_name)

    # Check if AI report draft is missing, generate it dynamically using Gemini
    if not inspection.ai_report_draft:
        try:
            from app.api.routes.ai import call_gemini
            
            prompt = f"""उत्तर प्रदेश सरकार के ग्राम विकास विभाग के मानकों के अनुसार एक अत्यंत औपचारिक और पेशेवर ग्राम पंचायत निरीक्षण रिपोर्ट (निरीक्षण आख्या) का हिंदी में मसौदा तैयार करें।

निरीक्षण विवरण:
- निरीक्षण संख्या (ID): {inspection.inspection_id}
- शीर्षक: {inspection.title}
- ग्राम पंचायत: {panchayat.name_hindi or panchayat.name if panchayat else 'N/A'} (जनपद: {inspection.district or (panchayat.district if panchayat else 'N/A')}, विकास खंड: {inspection.block or (panchayat.block if panchayat else 'N/A')})
- जांचकर्ता अधिकारी: {inspection.investigator_name or (engineer.name_hindi or engineer.name if engineer else 'N/A')} (पद: {engineer.designation or 'अवर अभियंता' if engineer else 'N/A'})
- परियोजना/कार्य का नाम: {inspection.project_name or 'N/A'} (कार्य कोड: {inspection.project_code or 'N/A'})
- निरीक्षण प्रकार: {inspection.inspection_type or 'सामान्य'}

मुख्य अवलोकन / कमियां (Observations):
{inspection.observations or 'स्थल निरीक्षण किया गया।'}

सुधार हेतु संस्तुतियां/सिफारिशें (Recommendations):
{inspection.recommendations or 'उचित सुधार कार्य किया जाए।'}

निम्नलिखित शीर्षकों के अंतर्गत पूर्ण हिंदी रिपोर्ट तैयार करें (भाषा विशुद्ध प्रशासनिक/सरकारी राजभाषा हिंदी होनी चाहिए):
1. कार्य का संक्षिप्त विवरण (Executive Summary)
2. निरीक्षण के दौरान पाई गई कमियां/विशिष्ट निष्कर्ष (Key Findings & Observations)
3. सुधार हेतु संस्तुतियां/सिफारिशें (Recommendations)
4. निष्कर्ष (Conclusion)

कृपया केवल हिंदी भाषा का उपयोग करें और सुनिश्चित करें कि भाषा का स्तर सरकारी पत्राचार और आधिकारिक आख्या के अनुरूप अत्यंत गरिमापूर्ण और गंभीर हो।"""
            
            ai_draft = await call_gemini(prompt, language="hi")
            if ai_draft and not ai_draft.startswith("AI Error:"):
                inspection.ai_report_draft = ai_draft
                await db.flush()
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"Failed to auto-generate Gemini report draft: {e}")

    try:
        build_pdf_report(inspection, panchayat, engineer, list(photos), list(approvals), output_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")

    # Save PDF report record
    file_size_kb = Path(output_path).stat().st_size // 1024
    report = Report(
        inspection_id=inspection_id,
        generated_by=current_user.id,
        file_path=output_path,
        file_name=file_name,
        file_size_kb=file_size_kb,
        report_format="pdf",
    )
    db.add(report)

    # DOCX Generation
    docx_file_name = f"Report_{inspection.inspection_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    docx_output_path = str(REPORTS_DIR / docx_file_name)
    try:
        build_docx_report(inspection, panchayat, engineer, list(photos), list(approvals), docx_output_path)
        docx_file_size_kb = Path(docx_output_path).stat().st_size // 1024
        docx_report = Report(
            inspection_id=inspection_id,
            generated_by=current_user.id,
            file_path=docx_output_path,
            file_name=docx_file_name,
            file_size_kb=docx_file_size_kb,
            report_format="docx",
        )
        db.add(docx_report)
    except Exception as docx_err:
        import logging
        logging.getLogger(__name__).error(f"DOCX generation failed: {docx_err}")

    return MessageResponse(
        message="Reports generated successfully",
        success=True,
        data={"file_name": file_name, "size_kb": file_size_kb},
    )


@router.get("/download/{inspection_id}")
async def download_report(
    inspection_id: str,
    format: str = "pdf",
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Download latest report PDF or DOCX for an inspection."""
    result = await db.execute(
        select(Report).where(Report.inspection_id == inspection_id)
        .where(Report.report_format == format)
        .order_by(Report.created_at.desc())
    )
    report = result.scalar_one_or_none()
    if not report or not Path(report.file_path).exists():
        raise HTTPException(status_code=404, detail=f"Report in {format} format not found. Generate it first.")

    media_type = "application/pdf" if format == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return FileResponse(
        report.file_path,
        media_type=media_type,
        filename=report.file_name,
    )

